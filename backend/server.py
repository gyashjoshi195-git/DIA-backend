from fastapi import FastAPI, APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
import os
import logging
import json
from pathlib import Path
from pydantic import BaseModel
from typing import Optional
import uuid
from datetime import datetime, timezone
from groq import AsyncGroq
import io

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

app = FastAPI()
api_router = APIRouter(prefix="/api")

GROQ_API_KEY = os.environ.get('GROQ_API_KEY')
groq_client = AsyncGroq(api_key=GROQ_API_KEY)

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


async def get_llm_response(system_message, user_text):
    response = await groq_client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[
            {"role": "system", "content": system_message},
            {"role": "user", "content": user_text}
        ]
    )
    return response.choices[0].message.content


# ─── Request Models ───
class GlobalOverviewRequest(BaseModel):
    topic: str

class CountryStanceRequest(BaseModel):
    country: str
    topic: str

class SpeechDrafterRequest(BaseModel):
    country: str
    agenda: str
    speech_time: int

class PositionPaperRequest(BaseModel):
    country: str
    committee: str
    agenda: str

class POIGeneratorRequest(BaseModel):
    target_country: str
    agenda: str

class DiplomaticShieldRequest(BaseModel):
    my_country: str
    topic: str
    poi_question: str

class FactCheckerRequest(BaseModel):
    claim: str

class ExportRequest(BaseModel):
    title: str
    content: str
    format: str

class ConferencePrepRequest(BaseModel):
    country: str
    committee: str
    agenda: str
    speech_time: int = 90


# ─── Endpoints ───
@api_router.get("/")
async def root():
    return {"message": "Diplomatic Intelligence Agent API"}


@api_router.post("/global-overview")
async def global_overview(req: GlobalOverviewRequest):
    system = """You are a senior international relations analyst and intelligence briefer for a Model United Nations conference.
KEY RULES:
- Combine emotional storytelling with rigorous statistical evidence
- Use the most impactful numbers for the specific agenda
- Only reference credible sources: UN (un.org), Council on Foreign Relations (cfr.org), Reuters (reuters.com), government sources (.gov), and National Intelligence Council (.nic)
- Include specific dollar amounts, percentages, resolution numbers, and treaty names
- Write in a formal but engaging intelligence briefing style"""

    user_msg = f"""Provide a comprehensive intelligence briefing on: "{req.topic}"

FORMAT YOUR RESPONSE EXACTLY AS FOLLOWS:

EXECUTIVE SUMMARY:
[Write exactly 3 paragraphs providing an executive summary. Include specific statistics, dollar figures, and data points. Reference real resolutions, treaties, or agreements by name/number.]

TOP NEWS & DEVELOPMENTS:
1. [Headline] | Source: [one of un.org, cfr.org, reuters.com, .gov, .nic] | [Brief 2-sentence summary with a specific data point]
2. [Headline] | Source: [source] | [Brief summary]
3. [Headline] | Source: [source] | [Brief summary]
4. [Headline] | Source: [source] | [Brief summary]
5. [Headline] | Source: [source] | [Brief summary]"""

    try:
        response = await get_llm_response(system, user_msg)
        return {"result": response, "topic": req.topic}
    except Exception as e:
        logger.error(f"Global overview error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@api_router.post("/country-stance")
async def country_stance(req: CountryStanceRequest):
    system = """You are a diplomatic intelligence analyst specializing in country position analysis for Model United Nations.
KEY RULES:
- Be specific about voting records (cite UN resolution numbers)
- Clearly distinguish between official positions and practical actions
- Red Lines should be concrete non-negotiables, not vague principles
- Allies and Adversaries must be contextual to the specific topic
- Include specific examples and historical precedents"""

    user_msg = f"""Analyze the diplomatic stance of {req.country} on the topic: "{req.topic}"

FORMAT YOUR RESPONSE EXACTLY AS:

OFFICIAL STANCE:
[2-3 paragraphs on the country's official position, including quotes from leaders and key policy documents]

VOTING RECORD:
[List 4-6 relevant UN resolutions with how this country voted: FOR, AGAINST, or ABSTAINED, with brief context]

RED LINES (Non-Negotiables):
1. [Specific red line with explanation]
2. [Specific red line with explanation]
3. [Specific red line with explanation]

ALLIES ON THIS ISSUE:
[List 4-5 allied countries with brief explanation of shared interests]

ADVERSARIES ON THIS ISSUE:
[List 4-5 opposing countries with brief explanation of conflicts]"""

    try:
        response = await get_llm_response(system, user_msg)
        return {"result": response, "country": req.country, "topic": req.topic}
    except Exception as e:
        logger.error(f"Country stance error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@api_router.post("/speech-drafter")
async def speech_drafter(req: SpeechDrafterRequest):
    word_count = int(req.speech_time * 2.5)

    system = f"""You are the world's best MUN speech writer. You use the "Maximum Humanizer" algorithm.

HUMANIZER RULES:
- Use rhetorical questions ("Are we really willing to stand by while...?")
- Vary sentence length: mix short punchy sentences with longer flowing ones
- NEVER use AI-isms: "It is important to note that...", "In conclusion...", "Furthermore...", "It is worth mentioning...", "As we navigate..."
- Include personal/emotional hooks ("When a child in Aleppo looks up at the sky...")
- Use "we" language to build consensus
- Weave 2-3 specific facts/numbers NATURALLY into the narrative (e.g., "$4.2 billion in trade", "15% increase in emissions", "Resolution 242")
- Numbers should feel like dramatic evidence, not data dumps
- The speech should build momentum: calm opening, rising tension, powerful close
- Target approximately {word_count} words for a {req.speech_time}-second speech
- Combine emotional storytelling with rigorous statistical evidence
- Search for the most impactful numbers for the specific agenda"""

    user_msg = f"""Write a MUN speech for the delegate of {req.country} on the agenda: "{req.agenda}"
Speech time limit: {req.speech_time} seconds (approximately {word_count} words).

The speech MUST include at least 2-3 specific facts, dollar figures, or resolution numbers woven naturally into the narrative.

DO NOT include any stage directions or labels like "[Opening]" or "(pause)". Just write the speech text itself."""

    try:
        response = await get_llm_response(system, user_msg)
        return {"result": response, "country": req.country, "agenda": req.agenda, "speech_time": req.speech_time, "word_count": len(response.split())}
    except Exception as e:
        logger.error(f"Speech drafter error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@api_router.post("/position-paper")
async def position_paper(req: PositionPaperRequest):
    system = """You are a formal diplomatic document writer specializing in MUN position papers. You write authoritative, well-structured position papers that demonstrate deep knowledge of both the issue and the country's stance.
Combine emotional storytelling with rigorous statistical evidence. Use the search results to find the most impactful numbers for the specific agenda."""

    user_msg = f"""Write a formal Position Paper for {req.country} in the {req.committee} committee on the agenda: "{req.agenda}"

FORMAT EXACTLY AS:

SECTION 1 - BACKGROUND OF THE ISSUE:
[3-4 paragraphs: Historical context, key events, current status, relevant UN resolutions with specific numbers and dates]

SECTION 2 - {req.country.upper()}'S NATIONAL POSITION & PAST ACTIONS:
[3-4 paragraphs: Official stance, domestic policies, international commitments, voting history, specific actions taken]

SECTION 3 - PROPOSED SOLUTIONS:
[Write 4-6 operative clauses in formal UN resolution format, starting with action verbs like "Calls upon...", "Urges...", "Recommends...", "Requests..."]

---

FACT CHECK & VERIFICATION:
[Review 4-5 major claims from the paper above and label each:]
- Claim: "[exact claim]" → [FACTUAL] / [MISLEADING] / [FALSE] - [Brief explanation]

POI TO EXPOSE INACCURACIES:
[For any MISLEADING or FALSE claims, write one aggressive Point of Information that could be used to challenge the delegate]"""

    try:
        response = await get_llm_response(system, user_msg)
        return {"result": response, "country": req.country, "committee": req.committee, "agenda": req.agenda}
    except Exception as e:
        logger.error(f"Position paper error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@api_router.post("/poi-generator")
async def poi_generator(req: POIGeneratorRequest):
    system = """You are a strategic debate advisor for Model United Nations. You specialize in identifying vulnerabilities in countries' positions and crafting devastating Points of Information (POIs).

KEY RULES:
- Base POIs on real controversies, human rights records, and policy contradictions
- Each POI should be framed as a question that exposes hypocrisy or weakness
- Include the specific evidence/context behind each POI
- POIs should be challenging but diplomatically phrased
- Rank from most devastating to least"""

    user_msg = f"""Generate the 5 best strategic POIs to ask the delegate of {req.target_country} on the agenda: "{req.agenda}"

FORMAT EACH POI EXACTLY AS:

POI 1 - [SEVERITY: HIGH/MEDIUM]:
Question: "[The actual POI question to ask]"
Evidence: [The real-world controversy, human rights record, or policy contradiction this is based on]
Expected Impact: [Why this question is effective]

POI 2 - [SEVERITY: HIGH/MEDIUM]:
...

POI 3 - [SEVERITY: HIGH/MEDIUM]:
...

POI 4 - [SEVERITY: MEDIUM/LOW]:
...

POI 5 - [SEVERITY: MEDIUM/LOW]:
..."""

    try:
        response = await get_llm_response(system, user_msg)
        return {"result": response, "target_country": req.target_country, "agenda": req.agenda}
    except Exception as e:
        logger.error(f"POI generator error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@api_router.post("/diplomatic-shield")
async def diplomatic_shield(req: DiplomaticShieldRequest):
    system = """You are a master MUN debater and diplomatic response coach. You specialize in "The Pivot" technique for handling hostile POIs.

THE PIVOT TECHNIQUE:
1. Brief Acknowledgment (1-2 sentences): Don't dodge the question entirely
2. The Pivot (core of the response): Shift focus to your country's strength OR a flaw in the opponent's position
3. The Close (1 sentence): End with a confident, slightly assertive statement

STYLE RULES:
- Sound confident and spontaneous, NOT scripted
- Use wit where appropriate (but stay diplomatic)
- Be slightly assertive — you're defending your nation
- Reference specific facts to back up your pivot
- Keep it under 30 seconds when spoken aloud (about 75 words)"""

    user_msg = f"""You are the delegate of {req.my_country}. The topic being discussed is: "{req.topic}"

Another delegate has asked you the following POI:
"{req.poi_question}"

Generate a "Pivot" response. Also provide a brief strategic analysis of why this response works.

FORMAT AS:

RESPONSE:
[The actual spoken response using The Pivot technique — confident, witty, assertive]

STRATEGY BREAKDOWN:
- Acknowledgment: [What you briefly conceded]
- The Pivot: [Where you shifted the focus]
- The Close: [Your power statement]
- Why It Works: [1-2 sentences on the rhetorical strategy]"""

    try:
        response = await get_llm_response(system, user_msg)
        return {"result": response, "my_country": req.my_country, "topic": req.topic}
    except Exception as e:
        logger.error(f"Diplomatic shield error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@api_router.post("/fact-checker")
async def fact_checker(req: FactCheckerRequest):
    system = """You are a ruthless diplomatic fact-checker for Model United Nations. You cross-reference claims against UN databases, Reuters, and .gov sources. You are blunt and aggressive when exposing falsehoods.

KEY RULES:
- Extract every significant factual claim from the input text
- Label each claim: [FACTUAL], [MISLEADING], or [FALSE]
- For [MISLEADING] and [FALSE] claims, generate one aggressive POI to expose the delegate
- Be specific with evidence — cite resolution numbers, dates, statistics
- Do NOT be generous — if a claim is even slightly misleading, call it out"""

    user_msg = f"""Analyze the following delegate's speech/claim and fact-check every significant statement:

"{req.claim}"

FORMAT YOUR RESPONSE EXACTLY AS:

CLAIM 1:
Statement: "[exact quote from the text]"
Verdict: [FACTUAL] or [MISLEADING] or [FALSE]
Evidence: [specific evidence with sources, dates, numbers]
POI: [aggressive POI question if MISLEADING or FALSE, or "N/A" if FACTUAL]

CLAIM 2:
Statement: "[exact quote]"
Verdict: [FACTUAL] or [MISLEADING] or [FALSE]
Evidence: [evidence]
POI: [POI or N/A]

(Continue for all significant claims found)

OVERALL ASSESSMENT:
Credibility Rating: [HIGH / MODERATE / LOW / DECEPTIVE]
Summary: [2-3 sentences on the overall accuracy and credibility of the speaker's statements]"""

    try:
        response = await get_llm_response(system, user_msg)
        return {"result": response, "claim": req.claim}
    except Exception as e:
        logger.error(f"Fact checker error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@api_router.post("/export")
async def export_document(req: ExportRequest):
    if req.format not in ("pdf", "docx"):
        raise HTTPException(status_code=400, detail="Format must be 'pdf' or 'docx'")
    try:
        if req.format == "pdf":
            from fpdf import FPDF
            pdf = FPDF()
            pdf.set_left_margin(20)
            pdf.set_right_margin(20)
            pdf.set_auto_page_break(auto=True, margin=20)
            pdf.add_page()
            safe_title = req.title.encode('latin-1', 'replace').decode('latin-1')
            pdf.set_font("Helvetica", "B", 14)
            pdf.multi_cell(0, 8, safe_title, new_x="LMARGIN", new_y="NEXT")
            pdf.ln(2)
            pdf.set_font("Helvetica", "I", 8)
            pdf.set_text_color(120, 120, 120)
            ts = datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')
            pdf.multi_cell(0, 5, f"Generated by Diplomatic Intelligence Agent | {ts}", new_x="LMARGIN", new_y="NEXT")
            pdf.ln(4)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("Helvetica", "", 10)
            safe_content = req.content.encode('latin-1', 'replace').decode('latin-1')
            headings = ('EXECUTIVE', 'OFFICIAL', 'VOTING', 'RED LINE', 'ALLIES', 'ADVERSARIES', 'SECTION', 'FACT CHECK', 'POI', 'RESPONSE', 'STRATEGY', 'TOP NEWS')
            for line in safe_content.split('\n'):
                stripped = line.strip()
                if not stripped:
                    pdf.ln(2)
                    continue
                if any(stripped.upper().startswith(h) for h in headings):
                    pdf.ln(3)
                    pdf.set_font("Helvetica", "B", 11)
                    pdf.multi_cell(0, 6, stripped, new_x="LMARGIN", new_y="NEXT")
                    pdf.set_font("Helvetica", "", 10)
                else:
                    pdf.multi_cell(0, 5, stripped, new_x="LMARGIN", new_y="NEXT")
            buffer = io.BytesIO()
            buffer.write(pdf.output())
            buffer.seek(0)
            fname = safe_title.replace(' ', '_')[:50]
            return StreamingResponse(buffer, media_type="application/pdf",
                headers={"Content-Disposition": f'attachment; filename="{fname}.pdf"'})

        elif req.format == "docx":
            from docx import Document as DocxDocument
            from docx.shared import Pt, Inches
            doc = DocxDocument()
            doc.add_heading(req.title, 0)
            doc.add_paragraph(f"Generated by Diplomatic Intelligence Agent | {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}").italic = True
            for line in req.content.split('\n'):
                stripped = line.strip()
                if not stripped:
                    continue
                if stripped.startswith(('EXECUTIVE', 'OFFICIAL', 'VOTING', 'RED LINE', 'ALLIES', 'ADVERSARIES', 'SECTION', 'FACT CHECK', 'POI', 'RESPONSE', 'STRATEGY', 'TOP NEWS')):
                    doc.add_heading(stripped, level=2)
                elif stripped.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '-')):
                    doc.add_paragraph(stripped, style='List Bullet')
                else:
                    doc.add_paragraph(stripped)
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            fname = req.title.replace(' ', '_')[:50]
            return StreamingResponse(buffer,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f'attachment; filename="{fname}.docx"'})
    except Exception as e:
        logger.error(f"Export error: {e}")
        raise HTTPException(status_code=500, detail=str(e))








@api_router.post("/conference-prep")
async def conference_prep(req: ConferencePrepRequest):
    async def generate():
        modules = [
            {
                "name": "global_overview",
                "label": "Global Overview",
                "handler": lambda: global_overview(GlobalOverviewRequest(topic=req.agenda))
            },
            {
                "name": "country_stance",
                "label": "Country Stance",
                "handler": lambda: country_stance(CountryStanceRequest(country=req.country, topic=req.agenda))
            },
            {
                "name": "speech_drafter",
                "label": "Speech Draft",
                "handler": lambda: speech_drafter(SpeechDrafterRequest(country=req.country, agenda=req.agenda, speech_time=req.speech_time))
            },
            {
                "name": "position_paper",
                "label": "Position Paper",
                "handler": lambda: position_paper(PositionPaperRequest(country=req.country, committee=req.committee, agenda=req.agenda))
            },
            {
                "name": "poi_defense",
                "label": "Anticipated POIs",
                "handler": lambda: poi_generator(POIGeneratorRequest(target_country=req.country, agenda=req.agenda))
            }
        ]

        for step, mod in enumerate(modules, 1):
            yield json.dumps({"step": step, "total": len(modules), "module": mod["name"], "label": mod["label"], "status": "generating"}) + "\n"
            try:
                result = await mod["handler"]()
                yield json.dumps({"step": step, "module": mod["name"], "label": mod["label"], "status": "complete", "data": result}) + "\n"
            except Exception as e:
                logger.error(f"Conference prep module {mod['name']} error: {e}")
                yield json.dumps({"step": step, "module": mod["name"], "label": mod["label"], "status": "error", "error": str(e)}) + "\n"

    return StreamingResponse(generate(), media_type="application/x-ndjson", headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})


app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)


